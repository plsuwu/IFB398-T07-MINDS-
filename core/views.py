# core/views.py
from __future__ import annotations
from pydoc import doc

from django.apps import apps
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.core.paginator import Paginator
from django.db.models import Q
from django.core.exceptions import PermissionDenied
from django.http import Http404, HttpResponse, JsonResponse, HttpResponseBadRequest
from django.urls import reverse
from django.views.decorators.http import require_GET, require_http_methods
from django.contrib import messages
from django.core.cache import cache
from django.shortcuts import render, get_object_or_404, redirect
from core.ai.report_service import generate_project_report
from .ai.granite_client import GraniteClient

from types import SimpleNamespace
import logging

# Exporting report
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
import re, io

from django.contrib.contenttypes.models import ContentType
from django.contrib.postgres.search import SearchQuery, SearchRank
from .forms import DocumentForm, DocumentSearchForm, ProspectForm, TenementForm, SampleForm, SurveyForm
from .models import Document, Process, SavedReport, AuditLog, log_audit, Prospect, DocLink, UserProfile, Drillhole, DrillholeSurvey, LithologyInterval, AssayResult, Organisation, Tenement, ApprovalWorkflow, Sample, Survey
from .importers import run_drillhole_import
from .permissions import role_required, clearance_required, log_view_access
from .utils import sha256_file, extract_text, chunk_text

from .tagging import TAG_LABEL


# ---------- Helpers ----------


def _get_model(app_label: str, model_name: str):
    """
    Best-effort dynamic model fetch (lets views work even if model doesn’t exist yet).
    """
    try:
        return apps.get_model(app_label, model_name)
    except LookupError:
        return None


def _count_model(app_label: str, model_name: str, where_clause: str = None) -> int:
    mdl = _get_model(app_label, model_name)
    if mdl is None:
        return 0
    return mdl.objects.count()


def _paginate(queryset, request, per_page: int = 20):
    paginator = Paginator(queryset, per_page)
    page_number = request.GET.get("page")
    return paginator.get_page(page_number)


def _org_qs_filter(request):
    """
    Returns a Q object for organisation-scoped queryset filtering.
    - Superusers: Q() — no restriction, see all data.
    - Authenticated users with an assigned organisation: Q(organisation=their_org).
    - Authenticated users with no organisation: Q(pk__in=[]) — see nothing.
    """
    if request.user.is_superuser:
        return Q()
    if request.user.is_authenticated and hasattr(request.user, 'profile'):
        org = request.user.profile.organisation
        if org is not None:
            return Q(organisation=org)
    return Q(pk__in=[])


# ---------- Landing / Dashboard ----------


@login_required
@require_GET
def home(request):
    """
    Simple landing that shows recent Projects & Documents (as per your snippet).
    """
    org_filter = _org_qs_filter(request)
    projects = Process.objects.filter(org_filter).order_by("-created_at")[:10]
    docs = Document.objects.filter(org_filter, is_latest=True).order_by("-created_at")[:10]
    return render(
        request,
        "core/home.html",
        {"projects": projects, "docs": docs},
    )


@login_required
@require_GET
def dashboard(request):
    """
    Dashboard cards + quick links. Works even if domain models aren’t ready yet.
    """
    org_filter = _org_qs_filter(request)
    Prospect = _get_model("core", "Prospect")
    Drillhole = _get_model("core", "Drillhole")
    Tenement = _get_model("core", "Tenement")
    metrics = {
        "project_count": Process.objects.filter(org_filter).count(),
        "document_count": Document.objects.filter(org_filter, is_latest=True).count(),
        "prospect_count": Prospect.objects.filter(org_filter).count() if Prospect else 0,
        "drillhole_count": Drillhole.objects.filter(org_filter).count() if Drillhole else 0,
        "tenement_count": Tenement.objects.filter(org_filter).count() if Tenement else 0,
    }
    recent_docs = Document.objects.filter(org_filter, is_latest=True).order_by("-created_at")[:8]
    return render(
        request,
        "core/dashboard.html",
        {"metrics": metrics, "recent_docs": recent_docs},
    )


# Optional: HTMX endpoint to refresh stats without reloading the whole page
@login_required
@require_GET
def stats_partial(request):
    org_filter = _org_qs_filter(request)
    Prospect = _get_model("core", "Prospect")
    Drillhole = _get_model("core", "Drillhole")
    Tenement = _get_model("core", "Tenement")
    ctx = {
        "project_count": Process.objects.filter(org_filter).count(),
        "document_count": Document.objects.filter(org_filter, is_latest=True).count(),
        "prospect_count": Prospect.objects.filter(org_filter).count() if Prospect else 0,
        "drillhole_count": Drillhole.objects.filter(org_filter).count() if Drillhole else 0,
        "tenement_count": Tenement.objects.filter(org_filter).count() if Tenement else 0,
    }
    return render(request, "core/partials/stats.html", ctx)


# ---------- Cache keys ----------

DOCS_CACHE_KEY = "docs:unfiltered:page1:v1"
DOCS_CACHE_TTL = 120  # 2 minutes


def _docs_cache_key(request):
    """Per-organisation cache key so users only see their own org's cached documents."""
    if request.user.is_superuser:
        return "docs:unfiltered:page1:v1:all"
    org = None
    if request.user.is_authenticated and hasattr(request.user, 'profile'):
        org = request.user.profile.organisation
    org_id = str(org.id) if org else "noorg"
    return f"docs:unfiltered:page1:v1:{org_id}"


# ---------- Documents ----------

log = logging.getLogger(__name__)


def _get_clearance_level(request) -> str:
    """return the requesting user's clearance level string, defaulting to PUBLIC"""
    if request.user.is_authenticated and hasattr(request.user, "profile"):
        return request.user.profile.clearance_level
    return "PUBLIC"


def _report_cache_key(process_id: str, clearance_level: str, latest_doc_ts) -> str:
    doc_fingerprint = latest_doc_ts.strftime("%Y%m%d%H%M%S%f") if latest_doc_ts else "empty"
    return f"report:v1:{process_id}:{clearance_level}:{doc_fingerprint}"


def _get_cached_report_bundle(process_id: str, clearance_level: str) -> dict:
    """
    Return cached {"md": str, "doc_ids": [str, ...]} for this process + clearance,
    generating and caching if not already present.

    Cache key includes latest-document timestamp so the cache auto-invalidates
    whenever a new document is added to the project.
    """
    latest_doc_ts = (
        Document.objects
        .filter(process_id=process_id)
        .order_by("-created_at")
        .values_list("created_at", flat=True)
        .first()
    )
    cache_key = _report_cache_key(process_id, clearance_level, latest_doc_ts)

    cached = cache.get(cache_key)
    # Regenerate for any non-dict cached value (None, legacy str, or tuple from old code)
    if not isinstance(cached, dict):
        md, doc_ids = generate_project_report(process_id, clearance_level=clearance_level)
        cached = {"md": md, "doc_ids": doc_ids}
        cache.set(cache_key, cached, 86400)  # 24 hours
    return cached


def _get_cached_report_md(process_id: str, clearance_level: str) -> str:
    return _get_cached_report_bundle(process_id, clearance_level)["md"]

@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.GEOLOGIST_MINE,
    UserProfile.RoleChoices.METALLURGIST,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["GET", "POST"])
def upload_doc(request):
    """
    Upload with SHA-256 de-duplication (your original logic, with tiny polish).
    """
    if request.method == "POST":
        _upload_org = getattr(getattr(request.user, "profile", None), "organisation", None)
        form = DocumentForm(request.POST, request.FILES, organisation=_upload_org)
        log.debug("FILES keys: %s", list(request.FILES.keys()))  # debug: ensure 'file' is present
        if form.is_valid():
            doc = form.save(commit=False)
            # Only set if user is authenticated (created_by is nullable)
            if request.user.is_authenticated:
                doc.created_by = request.user
            
            # Give extracted text a safe default in case extraction fails, to avoid null issues in search
            doc.extracted_text = ""

            if doc.file:
                # Important: call sha256_file on the uploaded file *before* saving
                doc.checksum_sha256 = sha256_file(doc.file)
                doc.extracted_text = extract_text(doc.file) or ""

            if doc.checksum_sha256 and Document.objects.filter(
                checksum_sha256=doc.checksum_sha256
            ).exists():
                # Duplicate detected — re-render with error + keep their form state
                docs = Document.objects.filter(_org_qs_filter(request), is_latest=True).order_by("-created_at")[:20]
                return render(
                    request,
                    "core/upload.html",
                    {
                        "form": form,
                        "docs": docs,
                        "error": "Duplicate file detected (checksum match).",
                    },
                )

            doc.extracted_text = doc.extracted_text or ""

            #Debug
            print("BEFORE SAVE extracted_text:", repr(doc.extracted_text))
            print("BEFORE SAVE type:", type(doc.extracted_text))
            print("BEFORE SAVE dict:", {
                "title": doc.title,
                "doc_type": doc.doc_type,
                "confidentiality": doc.confidentiality,
                "organisation_id": doc.organisation_id,
                "process_id": doc.process_id,
                "created_by_id": doc.created_by_id,
                "extracted_text": repr(doc.extracted_text),
            })

            doc.save()
            log_audit(request.user, AuditLog.ActionType.CREATE, doc,
                      f"Uploaded document '{doc.title}'",
                      ip_address=request.META.get("REMOTE_ADDR"))
            # Build text chunks for RAG retrieval
            if doc.extracted_text:
                from .models import DocumentChunk
                chunks = chunk_text(doc.extracted_text)
                DocumentChunk.objects.bulk_create([
                    DocumentChunk(
                    document=doc,
                    chunk_index=i,
                    text=chunk,
                    process=doc.process,
                    doc_type=doc.doc_type,
                    timestamp=doc.timestamp,
                    )
                    for i, chunk in enumerate(chunks)
                ])

            # Invalidate the unfiltered document list cache so the new doc appears immediately
            cache.delete(_docs_cache_key(request))

            # Pre-warm the report cache for this project (shifts LLM wait to upload time)
            if doc.process:
                uploader_clearance = _get_clearance_level(request)
                warm_cache_key = _report_cache_key(
                    str(doc.process_id), uploader_clearance, doc.created_at
                )
                try:
                    warm_md, warm_doc_ids = generate_project_report(
                        str(doc.process_id), clearance_level=uploader_clearance
                    )
                    cache.set(warm_cache_key, {"md": warm_md, "doc_ids": warm_doc_ids}, 86400)
                except Exception:
                    # Granite unavailable — report will be generated on first view request
                    pass

            return redirect("upload")
        else:
            # Show validation errors + keep the recent docs list
            # Show *why* it failed
            log.warning("Upload invalid: %s", form.errors)
            docs = Document.objects.filter(_org_qs_filter(request), is_latest=True).order_by("-created_at")[:20]
            return render(
                request,
                "core/upload.html",
                {"form": form, "docs": docs, "error": "Please correct the errors below."},
            )

    # GET
    _upload_org = getattr(getattr(request.user, "profile", None), "organisation", None)
    form = DocumentForm(organisation=_upload_org)
    docs = Document.objects.filter(_org_qs_filter(request), is_latest=True).order_by("-created_at")[:20]
    return render(request, "core/upload.html", {"form": form, "docs": docs})


@login_required
@require_GET
def documents(request):
    """
    Document library with search + tag + pagination.
    """
    # Build doc_type choices from whatever is actually in the DB
    existing_types = (
        Document.objects
        .filter(_org_qs_filter(request))
        .exclude(doc_type="")
        .exclude(doc_type__isnull=True)
        .values_list("doc_type", flat=True)
        .distinct()
        .order_by("doc_type")
    )

    type_choices = [("", "All types")] + [(t, t) for t in existing_types]

    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    form = DocumentSearchForm(request.GET or None, doc_type_choices=type_choices, organisation=org)
    qs = Document.objects.filter(_org_qs_filter(request), is_latest=True).select_related("process", "organisation").order_by("-created_at")

    q_value = ""
    if form.is_valid():

        # Full-text : title, doc_type, confidentiality, project name, org
        q = form.cleaned_data.get("q", "").strip()

        if q:
            search_query = SearchQuery(q, search_type='websearch', config='english')
            qs = (
                qs
                .annotate(rank=SearchRank('search_tsv', search_query))
                .filter(
                    Q(search_tsv=search_query)
                    | Q(doc_type__icontains=q)
                    | Q(confidentiality__icontains=q)
                    | Q(process__name__icontains=q)
                    | Q(organisation__name__icontains=q)
                )
                .order_by('-rank', '-created_at')
            )

        # Project
        process = form.cleaned_data.get("process")
        if process:
            qs = qs.filter(process=process)
 
        # Date range (inclusive, on the document's own date)
        date_from = form.cleaned_data.get("date_from")
        if date_from:
            qs = qs.filter(timestamp__gte=date_from)
 
        date_to = form.cleaned_data.get("date_to")
        if date_to:
            qs = qs.filter(timestamp__lte=date_to)
 
        # Metadata
        doc_type = form.cleaned_data.get("doc_type")
        if doc_type:
            qs = qs.filter(doc_type__iexact=doc_type)
 
        confidentiality = form.cleaned_data.get("confidentiality")
        if confidentiality:
            qs = qs.filter(confidentiality__iexact=confidentiality)
 
        tag = form.cleaned_data.get("tag")
        if tag:
            try:
                qs = qs.filter(tags__contains=[int(tag)])
            except (TypeError, ValueError):
                pass

        if form.cleaned_data.get("tenement"):
            qs = qs.filter(tenement=form.cleaned_data["tenement"])
        if form.cleaned_data.get("commodity"):
            qs = qs.filter(commodity__icontains=form.cleaned_data["commodity"])
        if form.cleaned_data.get("author_name"):
            qs = qs.filter(author_name__icontains=form.cleaned_data["author_name"])
        if form.cleaned_data.get("reporting_stage"):
            qs = qs.filter(reporting_stage=form.cleaned_data["reporting_stage"])

    filters_active = any(request.GET.get(f) for f in
                         ["q", "process", "date_from", "date_to", "doc_type", "confidentiality", "tag",
                          "tenement", "author_name", "commodity", "reporting_stage"])

    page_num = request.GET.get("page", "1")

    # Serve from cache for the default view (no filters, page 1)
    if not filters_active and page_num == "1":
        cached = cache.get(_docs_cache_key(request))
        if cached is not None:
            # Reconstruct a Page-like proxy from the cached dict so the template
            # interface (page.object_list, page.has_next, etc.) works unchanged.
            page_proxy = SimpleNamespace(
                object_list=cached["docs"],
                has_other_pages=cached["has_next"] or cached["has_previous"],
                has_previous=cached["has_previous"],
                has_next=cached["has_next"],
                number=1,
                paginator=SimpleNamespace(num_pages=cached["num_pages"]),
                previous_page_number=cached["prev_page_number"],
                next_page_number=cached["next_page_number"],
            )
            return render(request, "core/documents.html", {
                "form": form,
                "page": page_proxy,
                "q": "",
                "filters_active": False,
            })

    page = _paginate(qs, request, per_page=24)

    # cache only the unfiltered page-1 result abd Store a plain dict rather than the Page object to avoid serialising the full queryset into Redis
    if not filters_active and page_num == "1":
        cache.set(_docs_cache_key(request), {
            "docs": list(page.object_list),
            "num_pages": page.paginator.num_pages,
            "has_next": page.has_next(),
            "has_previous": page.has_previous(),
            "next_page_number": page.next_page_number() if page.has_next() else None,
            "prev_page_number": page.previous_page_number() if page.has_previous() else None,
        }, DOCS_CACHE_TTL)

    return render(request, "core/documents.html", {
        "form": form,
        "page": page,
        "q": q_value,
        "filters_active": filters_active,
    })


@login_required
@log_view_access(Document)
@require_GET
def document_detail(request, pk):
    doc = get_object_or_404(Document, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and doc.organisation
            and doc.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    tag_labels = [TAG_LABEL.get(t, f"Tag {t}") for t in (doc.tags or [])]
    version_family = doc.get_version_family()
    latest_version = version_family[-1] if version_family else doc
    profile = getattr(request.user, "profile", None)
    can_upload_version = profile and profile.role in (
        UserProfile.RoleChoices.FIELD_LEAD,
        UserProfile.RoleChoices.DATA_MANAGER,
        UserProfile.RoleChoices.ADMIN,
    )
    return render(request, "core/document_detail.html", {
        "doc": doc,
        "tag_labels": tag_labels,
        "version_family": version_family,
        "latest_version": latest_version,
        "can_upload_version": can_upload_version,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["POST", "DELETE"])
def delete_document(request, pk):
    """
    Delete a document and its associated file from storage (MinIO).
    """
    doc = get_object_or_404(Document, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and doc.organisation
            and doc.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    # Store title for success message
    doc_title = doc.title

    try:
        log_audit(request.user, AuditLog.ActionType.DELETE, doc,
                  f"Deleted document '{doc_title}'",
                  ip_address=request.META.get("REMOTE_ADDR"))
        doc.delete()

        # Invalidate the document list cache so the deletion is reflected immediately
        cache.delete(_docs_cache_key(request))

        # Return JSON response for HTMX/AJAX requests
        if request.headers.get('HX-Request'):
            return JsonResponse({
                "success": True,
                "message": f"Document '{doc_title}' deleted successfully."
            })

        # Redirect for regular form submissions
        return redirect("upload")

    except Exception as e:
        if request.headers.get('HX-Request'):
            return JsonResponse({
                "success": False,
                "message": f"Error deleting document: {str(e)}"
            }, status=500)

        # For regular requests, redirect with error (would need messages framework)
        return redirect("upload")


@login_required
@require_GET
def download_document(request, pk):
    doc = get_object_or_404(Document, pk=pk)
    profile = getattr(request.user, "profile", None)
    if not request.user.is_superuser:
        if profile and not profile.can_access_document(doc):
            raise PermissionDenied
    log_audit(
        request.user, AuditLog.ActionType.DOWNLOAD, doc,
        f"Downloaded '{doc.title}'",
        ip_address=request.META.get("REMOTE_ADDR"),
        user_agent=request.META.get("HTTP_USER_AGENT", "")[:500],
    )
    url = doc.file.url
    return redirect(url)


@login_required
@require_POST
@role_required(
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
def replace_document(request, pk):
    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    parent = get_object_or_404(Document, pk=pk, organisation=org)
    new_file = request.FILES.get("file")
    if not new_file:
        messages.error(request, "No file uploaded.")
        return redirect("document_detail", pk=pk)
    try:
        doc = Document.create_version(parent, new_file, request.user)
    except ValueError as e:
        messages.error(request, str(e))
        return redirect("document_detail", pk=pk)
    log_audit(
        request.user, AuditLog.ActionType.CREATE, doc,
        f"Uploaded version {doc.version_number} of '{parent.title}'",
        ip_address=request.META.get("REMOTE_ADDR"),
    )
    cache.delete(_docs_cache_key(request))
    messages.success(request, f"Version {doc.version_number} uploaded successfully.")
    return redirect("document_detail", pk=doc.pk)


# ---------- Projects / Domain pages (safe even if models are missing) ----------


@login_required
@require_GET
def projects(request):
    from django.db.models import Count
    org_filter = _org_qs_filter(request)
    qs = (
        Process.objects
        .filter(org_filter)
        .select_related("organisation")
        .annotate(
            prospect_count=Count("prospect", distinct=True),
            drillhole_count=Count("drillhole", distinct=True),
            document_count=Count("document", distinct=True),
        )
        .order_by("-created_at")
    )
    page = _paginate(qs, request)
    return render(request, "core/projects.html", {"page": page})


@login_required
@require_GET
def project_detail(request, pk):
    from .models import Tenement
    org_filter = _org_qs_filter(request)
    process = get_object_or_404(Process.objects.filter(org_filter), pk=pk)

    prospects_qs = Prospect.objects.filter(process=process).order_by("-created_at")
    drillholes_qs = Drillhole.objects.filter(process=process).order_by("name")
    tenements_qs = Tenement.objects.filter(process=process).order_by("name")
    documents_qs = Document.objects.filter(process=process, is_latest=True).order_by("-created_at")
    reports_qs = SavedReport.objects.filter(process=process).order_by("-created_at")

    return render(request, "core/project_detail.html", {
        "process":    process,
        "prospects":  prospects_qs,
        "drillholes": drillholes_qs,
        "tenements":  tenements_qs,
        "documents":  documents_qs[:10],
        "documents_total": documents_qs.count(),
        "reports":    reports_qs,
    })


@login_required
@require_GET
def prospects(request):
    Prospect = _get_model("core", "Prospect")
    if Prospect:
        qs = Prospect.objects.filter(_org_qs_filter(request)).order_by("-created_at")
        page = _paginate(qs, request)
    else:
        qs, page = [], None
    return render(
        request,
        "core/prospects.html",
        {"page": page, "model_exists": Prospect is not None},
    )


@login_required
def prospect_detail(request, pk):
    import json
    from django.core.serializers import serialize

    prospect = get_object_or_404(Prospect, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and prospect.organisation
            and prospect.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    doc_links = DocLink.objects.filter(
        content_type=ContentType.objects.get_for_model(Prospect),
        object_id=prospect.pk,
    ).select_related("document", "created_by").order_by("-created_at")

    drillholes = Drillhole.objects.filter(prospect=prospect).order_by("name")
    tenements = Tenement.objects.filter(process=prospect.process).order_by("name")
    prospect_reports = SavedReport.objects.filter(prospect=prospect).order_by("-created_at")
    samples = Sample.objects.filter(prospect=prospect).order_by("-created_at")
    surveys = Survey.objects.filter(prospect=prospect).order_by("-created_at")

    drillholes_geojson = serialize(
        'geojson',
        drillholes.exclude(collar_location__isnull=True),
        geometry_field='collar_location',
        fields=['name', 'depth', 'drill_type'],
    )
    tenements_geojson = serialize(
        'geojson',
        tenements.exclude(geom__isnull=True),
        geometry_field='geom',
        fields=['name'],
    )
    area_geom_geojson = json.dumps(json.loads(prospect.area_geom.geojson)) if prospect.area_geom else "null"

    return render(request, "core/prospect_detail.html", {
        "prospect":           prospect,
        "doc_links":          doc_links,
        "drillholes":         drillholes,
        "tenements":          tenements,
        "prospect_reports":   prospect_reports,
        "samples":            samples,
        "surveys":            surveys,
        "drillholes_geojson": drillholes_geojson,
        "tenements_geojson":  tenements_geojson,
        "area_geom_geojson":  area_geom_geojson,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["GET", "POST"])
def create_prospect(request):
    from django.contrib.gis.geos import Point

    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    if org is None:
        messages.error(request, "Your account is not linked to an organisation. Ask an administrator to assign one before creating prospects.")
        return redirect("prospects")

    initial_process_id = request.GET.get("project")
    initial_process = None
    if initial_process_id:
        try:
            initial_process = Process.objects.get(pk=initial_process_id, organisation=org)
        except Process.DoesNotExist:
            pass

    if request.method == "POST":
        form = ProspectForm(request.POST, organisation=org, initial_process=initial_process)
        if form.is_valid():
            prospect = form.save(commit=False)
            prospect.organisation = org
            lat = form.cleaned_data["latitude"]
            lng = form.cleaned_data["longitude"]
            prospect.geom = Point(float(lng), float(lat), srid=4326)
            prospect.save()
            log_audit(request.user, AuditLog.ActionType.CREATE, prospect,
                      f"Created prospect '{prospect.name}'",
                      ip_address=request.META.get("REMOTE_ADDR"))
            return redirect("prospect_detail", pk=prospect.pk)
        if request.headers.get("HX-Request"):
            return render(request, "core/partials/prospect_form_partial.html", {"form": form})
        return render(request, "core/prospect_form.html", {
            "form": form,
            "initial_lat": -25.0,
            "initial_lng": 133.0,
            "initial_zoom": 4,
            "initial_area_geojson": "null",
        })

    form = ProspectForm(organisation=org, initial_process=initial_process)
    return render(request, "core/prospect_form.html", {
        "form": form,
        "initial_lat": -25.0,
        "initial_lng": 133.0,
        "initial_zoom": 4,
        "initial_area_geojson": "null",
    })


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["GET", "POST"])
def edit_prospect(request, pk):
    from django.contrib.gis.geos import Point

    prospect = get_object_or_404(Prospect, pk=pk)
    org = getattr(getattr(request.user, "profile", None), "organisation", None)

    if not request.user.is_superuser and prospect.organisation != org:
        raise PermissionDenied

    if request.method == "POST":
        form = ProspectForm(request.POST, instance=prospect, organisation=org)
        if form.is_valid():
            updated = form.save(commit=False)
            lat = form.cleaned_data["latitude"]
            lng = form.cleaned_data["longitude"]
            updated.geom = Point(float(lng), float(lat), srid=4326)
            updated.save()
            log_audit(request.user, AuditLog.ActionType.EDIT, updated,
                      f"Edited prospect '{updated.name}'",
                      ip_address=request.META.get("REMOTE_ADDR"))
            return redirect("prospect_detail", pk=prospect.pk)
        import json as _json
        err_lat = prospect.geom.y if prospect.geom else -25.0
        err_lng = prospect.geom.x if prospect.geom else 133.0
        err_area = _json.dumps(_json.loads(prospect.area_geom.geojson)) if prospect.area_geom else "null"
        return render(request, "core/prospect_form.html", {
            "form": form,
            "editing": True,
            "prospect": prospect,
            "initial_lat": err_lat,
            "initial_lng": err_lng,
            "initial_zoom": 10 if prospect.geom else 4,
            "initial_area_geojson": err_area,
        })

    import json as _json
    initial_lat = prospect.geom.y if prospect.geom else -25.0
    initial_lng = prospect.geom.x if prospect.geom else 133.0
    initial_area = _json.dumps(_json.loads(prospect.area_geom.geojson)) if prospect.area_geom else "null"
    form = ProspectForm(instance=prospect, organisation=org)
    return render(request, "core/prospect_form.html", {
        "form": form,
        "editing": True,
        "prospect": prospect,
        "initial_lat": initial_lat,
        "initial_lng": initial_lng,
        "initial_zoom": 10 if prospect.geom else 4,
        "initial_area_geojson": initial_area,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_POST
def generate_prospect_report(request, pk):
    import hashlib
    prospect = get_object_or_404(Prospect, pk=pk)
    org_filter = _org_qs_filter(request)
    if not Prospect.objects.filter(org_filter, pk=pk).exists():
        raise PermissionDenied

    clearance_level = _get_clearance_level(request)
    report_title = request.POST.get("report_title", "").strip() or f"{prospect.name} — Prospect Report"

    try:
        md, doc_ids = generate_project_report(str(prospect.process_id), clearance_level=clearance_level)
    except Exception as e:
        log.error("Prospect report generation failed: %s", e)
        messages.error(request, f"Report generation failed: {e}")
        return redirect("prospect_detail", pk=pk)

    existing = SavedReport.objects.filter(
        prospect=prospect, title=report_title
    ).order_by("-version_number").first()

    if existing:
        report = SavedReport.create_version(
            parent=existing,
            content_md=md,
            user=request.user,
            reason=SavedReport.ChangeReason.REGENERATED,
        )
    else:
        report = SavedReport.objects.create(
            process=prospect.process,
            organisation=prospect.organisation,
            prospect=prospect,
            title=report_title,
            content_md=md,
            content_hash=hashlib.sha256(md.encode()).hexdigest(),
            created_by=request.user,
            version_number=1,
            change_reason=SavedReport.ChangeReason.GENERATED,
        )

    if doc_ids:
        report.source_documents.set(Document.objects.filter(pk__in=doc_ids))

    return redirect("saved_report_editor", report_id=report.pk)


# ---------- Prospect–Report Assignment ----------


@login_required
@require_POST
def assign_report_prospect(request, report_id):
    """Assign or clear the prospect linked to a saved report."""
    report = get_object_or_404(SavedReport, pk=report_id)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and report.organisation
            and report.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    prospect_id = request.POST.get("prospect_id") or None
    if prospect_id:
        report.prospect = Prospect.objects.filter(
            pk=prospect_id, process=report.process
        ).first()
    else:
        report.prospect = None
    report.save(update_fields=["prospect"])
    return JsonResponse({
        "success": True,
        "prospect_id": str(report.prospect_id) if report.prospect_id else None,
    })


# ---------- Samples ----------


@login_required
@require_GET
def samples(request):
    qs = Sample.objects.filter(_org_qs_filter(request)).select_related("process", "prospect").order_by("-created_at")
    prospect_id = request.GET.get("prospect")
    if prospect_id:
        qs = qs.filter(prospect_id=prospect_id)
    page = _paginate(qs, request)
    return render(request, "core/samples.html", {"page": page})


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["GET", "POST"])
def create_sample(request):
    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    if org is None:
        messages.error(request, "Your account is not linked to an organisation.")
        return redirect("samples")

    initial_prospect = None
    prospect_id = request.GET.get("prospect")
    if prospect_id:
        initial_prospect = Prospect.objects.filter(pk=prospect_id, organisation=org).first()

    if request.method == "POST":
        form = SampleForm(request.POST, organisation=org)
        if form.is_valid():
            sample = form.save(commit=False)
            sample.organisation = org
            sample.save()
            log_audit(request.user, AuditLog.ActionType.CREATE, sample,
                      f"Created sample '{sample.name}'",
                      ip_address=request.META.get("REMOTE_ADDR"))
            redirect_to = request.POST.get("next") or "sample_detail"
            if redirect_to == "prospect" and sample.prospect_id:
                return redirect("prospect_detail", pk=sample.prospect_id)
            return redirect("sample_detail", pk=sample.pk)
        return render(request, "core/sample_form.html", {"form": form})

    form = SampleForm(organisation=org, initial_prospect=initial_prospect)
    return render(request, "core/sample_form.html", {"form": form, "initial_prospect": initial_prospect})


@login_required
@require_GET
def sample_detail(request, pk):
    sample = get_object_or_404(Sample, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and sample.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    return render(request, "core/sample_detail.html", {"sample": sample})


# ---------- Surveys ----------


@login_required
@require_GET
def surveys(request):
    qs = Survey.objects.filter(_org_qs_filter(request)).select_related("process", "prospect").order_by("-created_at")
    prospect_id = request.GET.get("prospect")
    if prospect_id:
        qs = qs.filter(prospect_id=prospect_id)
    page = _paginate(qs, request)
    return render(request, "core/surveys.html", {"page": page})


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["GET", "POST"])
def create_survey(request):
    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    if org is None:
        messages.error(request, "Your account is not linked to an organisation.")
        return redirect("surveys")

    initial_prospect = None
    prospect_id = request.GET.get("prospect")
    if prospect_id:
        initial_prospect = Prospect.objects.filter(pk=prospect_id, organisation=org).first()

    if request.method == "POST":
        form = SurveyForm(request.POST, organisation=org)
        if form.is_valid():
            survey = form.save(commit=False)
            survey.organisation = org
            survey.save()
            log_audit(request.user, AuditLog.ActionType.CREATE, survey,
                      f"Created survey '{survey.name}'",
                      ip_address=request.META.get("REMOTE_ADDR"))
            if survey.prospect_id:
                return redirect("prospect_detail", pk=survey.prospect_id)
            return redirect("survey_detail", pk=survey.pk)
        return render(request, "core/survey_form.html", {"form": form})

    form = SurveyForm(organisation=org, initial_prospect=initial_prospect)
    return render(request, "core/survey_form.html", {"form": form, "initial_prospect": initial_prospect})


@login_required
@require_GET
def survey_detail(request, pk):
    survey = get_object_or_404(Survey, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and survey.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    area_geom_geojson = "null"
    if survey.geom:
        import json
        area_geom_geojson = json.dumps(json.loads(survey.geom.geojson))
    return render(request, "core/survey_detail.html", {
        "survey": survey,
        "area_geom_geojson": area_geom_geojson,
    })


# ---------- DocLink Views ----------

_LINKABLE_MODELS = {
    "prospect": ("core", "prospect"),
    "tenement": ("core", "tenement"),
    "drillhole": ("core", "drillhole"),
    "process": ("core", "process"),
}


@login_required
@require_http_methods(["GET"])
def doc_link_picker(request):
    """HTMX partial: render the document picker modal for linking a document to an entity."""
    content_type_label = request.GET.get("content_type", "")
    object_id = request.GET.get("object_id", "")

    if content_type_label not in _LINKABLE_MODELS:
        return HttpResponseBadRequest("Invalid content type.")

    documents = Document.objects.filter(_org_qs_filter(request)).order_by("-created_at")[:100]
    return render(request, "core/partials/doc_link_picker.html", {
        "documents": documents,
        "content_type_label": content_type_label,
        "object_id": object_id,
    })


@login_required
@require_POST
def create_doc_link(request):
    """Create one or more DocLinks between documents and a target entity."""
    document_ids = request.POST.getlist("document_id")
    content_type_label = request.POST.get("content_type_label")
    object_id = request.POST.get("object_id")

    if not document_ids or not content_type_label or not object_id:
        return HttpResponseBadRequest("Missing required fields.")

    if content_type_label not in _LINKABLE_MODELS:
        return HttpResponseBadRequest("Invalid content type.")

    app_label, model_name = _LINKABLE_MODELS[content_type_label]
    try:
        ct = ContentType.objects.get(app_label=app_label, model=model_name)
    except ContentType.DoesNotExist:
        return HttpResponseBadRequest("Content type not found.")

    created_by = request.user if request.user.is_authenticated else None
    for document_id in document_ids:
        document = get_object_or_404(Document, pk=document_id)
        link, created = DocLink.objects.get_or_create(
            document=document,
            content_type=ct,
            object_id=object_id,
            defaults={"created_by": created_by},
        )
        if created:
            log_audit(request.user, AuditLog.ActionType.EDIT, document,
                      f"Linked document '{document.title}' to {content_type_label} {object_id}",
                      ip_address=request.META.get("REMOTE_ADDR"))

    if content_type_label == "prospect":
        prospect = get_object_or_404(Prospect, pk=object_id)
        doc_links = DocLink.objects.filter(
            content_type=ct,
            object_id=object_id,
        ).select_related("document", "created_by").order_by("-created_at")
        return render(request, "core/partials/linked_documents.html", {
            "entity": prospect,
            "doc_links": doc_links,
            "content_type_label": content_type_label,
        })

    return HttpResponse(status=204)


@login_required
@require_POST
def delete_doc_link(request, pk):
    """Delete a DocLink record and re-render the linked documents section."""
    link = get_object_or_404(DocLink, pk=pk)
    ct = link.content_type
    object_id = link.object_id
    content_type_label = ct.model

    log_audit(request.user, AuditLog.ActionType.DELETE, link.document,
              f"Unlinked document '{link.document.title}' from {content_type_label} {object_id}",
              ip_address=request.META.get("REMOTE_ADDR"))
    link.delete()

    if content_type_label == "prospect":
        prospect = get_object_or_404(Prospect, pk=object_id)
        doc_links = DocLink.objects.filter(
            content_type=ct,
            object_id=object_id,
        ).select_related("document", "created_by").order_by("-created_at")
        return render(request, "core/partials/linked_documents.html", {
            "entity": prospect,
            "doc_links": doc_links,
            "content_type_label": content_type_label,
        })

    return HttpResponse(status=204)


# ---------- Drillhole–Prospect Linking ----------


@login_required
@require_GET
def drillhole_link_picker(request):
    prospect_id = request.GET.get("prospect_id")
    prospect = get_object_or_404(Prospect, pk=prospect_id)
    available_drillholes = Drillhole.objects.filter(
        process=prospect.process,
        prospect__isnull=True,
    ).order_by("name")
    return render(request, "core/partials/drillhole_link_picker.html", {
        "available_drillholes": available_drillholes,
        "prospect_id": prospect_id,
    })


@login_required
@require_POST
def link_drillhole(request):
    drillhole_id = request.POST.get("drillhole_id")
    prospect_id  = request.POST.get("prospect_id")
    prospect  = get_object_or_404(Prospect, pk=prospect_id)
    drillhole = get_object_or_404(Drillhole, pk=drillhole_id)
    drillhole.prospect = prospect
    drillhole.save(update_fields=["prospect"])
    drillholes = Drillhole.objects.filter(prospect=prospect).order_by("name")
    return render(request, "core/partials/linked_drillholes.html", {
        "prospect":   prospect,
        "drillholes": drillholes,
    })


@login_required
@require_POST
def unlink_drillhole(request, pk):
    drillhole = get_object_or_404(Drillhole, pk=pk)
    prospect  = drillhole.prospect
    drillhole.prospect = None
    drillhole.save(update_fields=["prospect"])
    drillholes = Drillhole.objects.filter(prospect=prospect).order_by("name")
    return render(request, "core/partials/linked_drillholes.html", {
        "prospect":   prospect,
        "drillholes": drillholes,
    })


@login_required
@require_POST
def bulk_link_drillholes(request):
    """HTMX: link multiple drillholes to a prospect in one action."""
    prospect_id  = request.POST.get("prospect_id")
    drillhole_ids = request.POST.getlist("drillhole_id")
    prospect = get_object_or_404(Prospect, pk=prospect_id)
    if drillhole_ids:
        Drillhole.objects.filter(
            pk__in=drillhole_ids,
            process=prospect.process,
        ).update(prospect=prospect)
    drillholes = Drillhole.objects.filter(prospect=prospect).order_by("name")
    return render(request, "core/partials/linked_drillholes.html", {
        "prospect":   prospect,
        "drillholes": drillholes,
    })


@login_required
@require_POST
def bulk_assign_drillholes(request):
    """Bulk-assign selected drillholes to a prospect from the drillholes list page."""
    prospect_id   = request.POST.get("prospect_id")
    drillhole_ids = request.POST.getlist("drillhole_ids")

    if not prospect_id or not drillhole_ids:
        messages.error(request, "Select a prospect and at least one drillhole.")
        return redirect("drillholes")

    prospect = get_object_or_404(Prospect, pk=prospect_id)
    org_filter = _org_qs_filter(request)
    if not Prospect.objects.filter(org_filter, pk=prospect_id).exists():
        raise PermissionDenied

    updated = Drillhole.objects.filter(
        pk__in=drillhole_ids,
        process=prospect.process,
    ).update(prospect=prospect)

    skipped = len(drillhole_ids) - updated
    msg = f"{updated} drillhole{'s' if updated != 1 else ''} linked to {prospect.name}."
    if skipped:
        msg += f" {skipped} skipped (different project)."
    messages.success(request, msg)
    return redirect("drillholes")


@login_required
@require_GET
def drillholes(request):
    Drillhole = _get_model("core", "Drillhole")
    if Drillhole:
        qs = Drillhole.objects.filter(_org_qs_filter(request)).order_by("-created_at")
        page = _paginate(qs, request)
    else:
        qs, page = [], None
    prospects = Prospect.objects.filter(_org_qs_filter(request)).select_related("process").order_by("process__name", "name")
    return render(
        request,
        "core/drillholes.html",
        {"page": page, "model_exists": Drillhole is not None, "prospects": prospects},
    )


@login_required
@require_GET
def drillhole_detail(request, pk):
    drillhole = get_object_or_404(Drillhole, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, "profile")
            and request.user.profile.organisation
            and drillhole.organisation
            and drillhole.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    surveys = DrillholeSurvey.objects.filter(drillhole=drillhole).order_by("depth")
    lithology = LithologyInterval.objects.filter(drillhole=drillhole).order_by("from_depth")
    assays = AssayResult.objects.filter(drillhole=drillhole).order_by("from_depth")
    return render(request, "core/drillhole_detail.html", {
        "drillhole": drillhole,
        "surveys":   surveys,
        "lithology": lithology,
        "assays":    assays,
    })


@login_required
@require_http_methods(["GET", "POST"])
def drillhole_import(request):
    if not request.user.is_superuser:
        profile = getattr(request.user, "profile", None)
        allowed = (UserProfile.RoleChoices.DATA_MANAGER, UserProfile.RoleChoices.ADMIN)
        if not profile or profile.role not in allowed:
            raise PermissionDenied

    org_filter = _org_qs_filter(request)
    organisations = Organisation.objects.filter(org_filter).order_by("name")
    processes     = Process.objects.filter(org_filter).select_related("organisation").order_by("name")

    context = {"organisations": organisations, "processes": processes}

    if request.method == "POST":
        file    = request.FILES.get("file")
        org_id  = request.POST.get("organisation", "").strip()
        proc_id = request.POST.get("process", "").strip()
        dry_run = request.POST.get("dry_run") == "on"
        update  = request.POST.get("update") == "on"

        form_errors = []
        if not file:
            form_errors.append("No file selected.")
        if not org_id:
            form_errors.append("No organisation selected.")
        if not proc_id:
            form_errors.append("No process selected.")

        if not form_errors:
            try:
                org     = Organisation.objects.get(pk=org_id)
                process = Process.objects.get(pk=proc_id, organisation=org)
                result  = run_drillhole_import(
                    file, org=org, process=process, dry_run=dry_run, update=update
                )
                if not dry_run:
                    c = result.get("counters", {})
                    log_audit(request.user, AuditLog.ActionType.CREATE, process,
                              f"Imported drillholes: {c.get('created', 0)} created, "
                              f"{c.get('updated', 0)} updated, {c.get('skipped', 0)} skipped",
                              ip_address=request.META.get("REMOTE_ADDR"))
                context.update({
                    "result":       result,
                    "counters":     result["counters"],
                    "import_errors": result["errors"],
                    "dry_run":      dry_run,
                    "selected_org":  org_id,
                    "selected_proc": proc_id,
                })
            except Exception as e:
                log.error("Drillhole import failed: %s", e, exc_info=True)
                form_errors.append(str(e))

        context["form_errors"] = form_errors

    return render(request, "core/drillhole_import.html", context)


@login_required
@require_GET
def tenements(request):
    Tenement = _get_model("core", "Tenement")
    if Tenement:
        qs = Tenement.objects.filter(_org_qs_filter(request)).order_by("-created_at")
        page = _paginate(qs, request)
    else:
        qs, page = [], None
    return render(
        request,
        "core/tenements.html",
        {"page": page, "model_exists": Tenement is not None},
    )


@login_required
@require_http_methods(["GET", "POST"])
def create_tenement(request):
    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    if request.method == "POST":
        form = TenementForm(request.POST, organisation=org)
        if form.is_valid():
            tenement = form.save()
            log_audit(request.user, AuditLog.ActionType.CREATE, tenement,
                      f"Created tenement '{tenement.name}'")
            messages.success(request, f"Tenement '{tenement.name}' created.")
            return redirect("tenement_detail", pk=tenement.pk)
    else:
        initial = {}
        process_pk = request.GET.get("process")
        if process_pk:
            initial["process"] = process_pk
        form = TenementForm(organisation=org, initial=initial)
    return render(request, "core/tenement_form.html", {"form": form, "editing": False})


@login_required
@require_GET
def tenement_detail(request, pk):
    tenement = get_object_or_404(
        Tenement.objects.filter(_org_qs_filter(request)), pk=pk
    )
    documents = (
        Document.objects.filter(tenement=tenement, is_latest=True)
        .order_by("-created_at")[:10]
    )
    return render(request, "core/tenement_detail.html", {
        "tenement": tenement,
        "documents": documents,
    })


@login_required
@require_http_methods(["GET", "POST"])
def edit_tenement(request, pk):
    tenement = get_object_or_404(
        Tenement.objects.filter(_org_qs_filter(request)), pk=pk
    )
    org = getattr(getattr(request.user, "profile", None), "organisation", None)
    if request.method == "POST":
        form = TenementForm(request.POST, instance=tenement, organisation=org)
        if form.is_valid():
            form.save()
            log_audit(request.user, AuditLog.ActionType.EDIT, tenement,
                      f"Edited tenement '{tenement.name}'")
            messages.success(request, f"Tenement '{tenement.name}' updated.")
            return redirect("tenement_detail", pk=tenement.pk)
    else:
        form = TenementForm(instance=tenement, organisation=org)
    initial_geojson = tenement.geom.json if tenement.geom else ""
    return render(request, "core/tenement_form.html", {
        "form": form,
        "editing": True,
        "tenement": tenement,
        "initial_geojson": initial_geojson,
    })


@login_required
@require_http_methods(["GET", "POST"])
def edit_process_geometry(request, pk):
    process = get_object_or_404(
        Process.objects.filter(_org_qs_filter(request)), pk=pk
    )
    if request.method == "POST":
        from django.contrib.gis.geos import GEOSGeometry, MultiPolygon, Polygon
        geojson = request.POST.get("geom_geojson", "").strip()
        if geojson:
            try:
                geom = GEOSGeometry(geojson)
                if isinstance(geom, Polygon):
                    geom = MultiPolygon(geom)
                process.geom = geom
                process.save()
                log_audit(request.user, AuditLog.ActionType.EDIT, process,
                          f"Updated spatial boundary for '{process.name}'")
                messages.success(request, "Project boundary updated.")
                return redirect("project_detail", pk=process.pk)
            except Exception:
                messages.error(request, "Invalid geometry — please redraw the boundary.")
        else:
            messages.error(request, "No geometry provided — please draw a boundary on the map.")

    initial_geojson = process.geom.json if process.geom else ""
    return render(request, "core/process_geometry_form.html", {
        "process": process,
        "initial_geojson": initial_geojson,
    })


# ---------- AI / Map / Utilities ----------


@login_required
@require_GET
def ai_insights(request):
    """
    Placeholder page for AI features (report generation, summarization, etc.).
    """
    # You can pass recent docs/projects for prompts, etc.
    org_filter = _org_qs_filter(request)
    return render(
        request,
        "core/ai_insights.html",
        {
            "recent_docs": Document.objects.filter(org_filter).order_by("-created_at")[:12],
            "recent_projects": Process.objects.filter(org_filter).order_by("-created_at")[:8],
            "recent_reports": SavedReport.objects.filter(org_filter).select_related("process").order_by("-created_at")[:10],
        },
    )


@login_required
@require_GET
def map_view(request):
    """
    Simple map page that we should consider wiring to Leaflet / PostGIS endpoints.
    """
    return render(request, "core/map.html")


@require_GET
def healthcheck(request):
    """
    Lightweight container health endpoint (used by k8s/docker healthchecks later).
    """
    return JsonResponse({"status": "ok"})




@login_required
@require_GET
def project_report_pdf(request, process_id: str):
    org_filter = _org_qs_filter(request)
    if not Process.objects.filter(org_filter, pk=process_id).exists():
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    try:
        md_text = _get_cached_report_md(process_id, clearance_level)
    except Exception as e:
        log.error("PDF export failed for process %s: %s", process_id, e)
        return HttpResponse("Report generation failed: Granite model unavailable.", status=503, content_type="text/plain")
    process = Process.objects.get(pk=process_id)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    # Custom styles
    h1 = ParagraphStyle('h1', parent=styles['Heading1'], textColor=colors.HexColor('#0e7490'), spaceAfter=10)
    h2 = ParagraphStyle('h2', parent=styles['Heading2'], textColor=colors.HexColor('#155e75'), spaceAfter=6)
    h3 = ParagraphStyle('h3', parent=styles['Heading3'], textColor=colors.HexColor('#1e4d5c'), spaceAfter=4)
    body = ParagraphStyle('body', parent=styles['Normal'], spaceAfter=6, leading=16)
    bullet = ParagraphStyle('bullet', parent=styles['Normal'], leftIndent=20, spaceAfter=4,
                             bulletIndent=10, leading=16)

    story = []
    for line in md_text.splitlines():
        if line.startswith('### '):
            story.append(Paragraph(line[4:], h3))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], h2))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], h1))
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph(f'• {line[2:]}', bullet))
        elif re.match(r'^\d+\. ', line):
            story.append(Paragraph(re.sub(r'^\d+\. ', '', line), bullet))
        elif line.strip() == '':
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(line, body))

    doc.build(story)
    buf.seek(0)
    slug = re.sub(r'[^\w-]', '_', process.name or str(process_id))
    response = HttpResponse(buf.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{slug}_report.pdf"'
    return response

@login_required
@require_GET
def project_report_docx(request, process_id: str):
    org_filter = _org_qs_filter(request)
    if not Process.objects.filter(org_filter, pk=process_id).exists():
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    try:
        md_text = _get_cached_report_md(process_id, clearance_level)
    except Exception as e:
        log.error("DOCX export failed for process %s: %s", process_id, e)
        return HttpResponse("Report generation failed: Granite model unavailable.", status=503, content_type="text/plain")
    process = Process.objects.get(pk=process_id)

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in md_text.splitlines():
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    slug = re.sub(r"[^\w-]", "_", process.name or str(process_id))
    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{slug}_report.docx"'
    return response

@login_required
def document_analysis_detail(request, pk):
    document = get_object_or_404(Document, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and document.organisation
            and document.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    analysis_text = getattr(document, "analysis_text", "") or "No insights available yet."

    return render(request, "core/document_analysis_detail.html", {
        "document": document,
        "analysis": analysis_text,
    })

@login_required
@require_POST
def save_document_analysis(request, pk):
    return HttpResponse("Save document analysis placeholder")


# ---------- GeoJSON API Endpoints for Map Viewer ----------


@login_required
@require_GET
def geojson_projects(request):
    """
    GeoJSON endpoint for Process (projects/operations) with spatial data
    Returns all processes with geometry for da map
    """
    from django.core.serializers import serialize
    from .models import Process

    # Only include processes with geometry
    processes = Process.objects.filter(
        _org_qs_filter(request), geom__isnull=False
    ).select_related('organisation')

    if not processes.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    # Use GeoDjangos built in serialiser (lets us take coordinates and translate into GeoJSOn text for geodata)
    geojson_data = serialize(
        'geojson',
        processes,
        geometry_field='geom',
        fields=('name', 'mode', 'commodity', 'organisation')
    )

    # Parse and return as JSON (serialise returns a string )
    import json
    return JsonResponse(json.loads(geojson_data), safe=False)


@login_required
@require_GET
def geojson_tenements(request):
    """
    GeoJSON endpoint for Tenement boundaries.
    Returns all tenements with geometry for map visualisation
    """
    from django.core.serializers import serialize
    from .models import Tenement

    tenements = Tenement.objects.filter(
        _org_qs_filter(request), geom__isnull=False
    ).select_related('organisation', 'process')

    if not tenements.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    geojson_data = serialize(
        'geojson',
        tenements,
        geometry_field='geom',
        fields=('name', 'organisation', 'process')
    )

    import json
    return JsonResponse(json.loads(geojson_data), safe=False)


@login_required
@require_GET
def geojson_prospects(request):
    """
    GeoJSON endpoint for Prospect locations.
    Includes area_geom (polygon) in properties when present.
    """
    import json

    prospects = Prospect.objects.filter(
        _org_qs_filter(request), geom__isnull=False
    ).select_related('organisation', 'process')

    features = []
    for p in prospects:
        props = {
            "pk": str(p.pk),
            "name": p.name,
            "organisation": str(p.organisation) if p.organisation else None,
            "process": str(p.process) if p.process else None,
        }
        if p.area_geom:
            props["area_geom_geojson"] = json.loads(p.area_geom.geojson)
        features.append({
            "type": "Feature",
            "geometry": json.loads(p.geom.geojson),
            "properties": props,
        })

    return JsonResponse({"type": "FeatureCollection", "features": features})


@login_required
@require_GET
def geojson_drillholes(request):
    """
    GeoJSON endpoint for Drillhole collar locations
    Returns all drillholes with collar locations 
    """
    from django.core.serializers import serialize
    from .models import Drillhole

    drillholes = Drillhole.objects.filter(
        _org_qs_filter(request), collar_location__isnull=False
    ).select_related('organisation', 'process')

    if not drillholes.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    geojson_data = serialize(
        'geojson',
        drillholes,
        geometry_field='collar_location',
        fields=('name', 'depth', 'azimuth', 'dip', 'organisation', 'process')
    )

    import json
    return JsonResponse(json.loads(geojson_data), safe=False)


@login_required
@require_POST
def spatial_search(request):
    """
    POST body (application/json):
        { "geometry": <GeoJSON geometry object> }
        Optionally with a radius for point searches:
        { "geometry": {"type": "Point", "coordinates": [lng, lat]}, "radius": 5000 }
        radius is in metres; internally converted to an approximate buffer in degrees.

    Returns an HTML partial (templates/core/partials/spatial_search_results.html)
    for injection into the map's results panel.
    """
    import json
    from django.contrib.gis.geos import GEOSGeometry

    try:
        body = json.loads(request.body)
        geometry_json = body.get("geometry")
        if not geometry_json:
            return HttpResponseBadRequest("No geometry provided.")
        geom = GEOSGeometry(json.dumps(geometry_json), srid=4326)
        radius_m = body.get("radius")
    except Exception as e:
        return HttpResponseBadRequest(f"Invalid geometry: {e}")

    org_filter = _org_qs_filter(request)

    # For a point + radius: buffer the point into an approximate circle
    # (1 degree ≈ 111,111 m at the equator; acceptable approximation for Australia)
    if geom.geom_type == 'Point' and radius_m:
        radius_deg = float(radius_m) / 111111.0
        search_geom = geom.buffer(radius_deg)
    else:
        search_geom = geom

    ProspectModel = _get_model("core", "Prospect")
    TenementModel = _get_model("core", "Tenement")
    DrillholeModel = _get_model("core", "Drillhole")

    processes  = Process.objects.filter(org_filter, geom__intersects=search_geom).order_by('name').values('id', 'name', 'mode')
    tenements  = TenementModel.objects.filter(org_filter, geom__intersects=search_geom).order_by('name').values('id', 'name') if TenementModel else []
    prospects  = ProspectModel.objects.filter(org_filter, geom__intersects=search_geom).order_by('name').values('id', 'name') if ProspectModel else []
    drillholes = DrillholeModel.objects.filter(org_filter, collar_location__intersects=search_geom).order_by('name').values('id', 'name', 'depth') if DrillholeModel else []

    # Find documents linked to matched spatial entities via DocLink (Group 1)
    documents = []
    DocLinkModel = _get_model("core", "DocLink")
    if DocLinkModel and (prospects or drillholes or processes):
        from django.contrib.contenttypes.models import ContentType
        linked_doc_ids = set()

        if ProspectModel and prospects:
            ct = ContentType.objects.get_for_model(ProspectModel)
            ids = [str(r['id']) for r in prospects]
            linked_doc_ids.update(
                DocLinkModel.objects.filter(content_type=ct, object_id__in=ids)
                .values_list('document_id', flat=True)
            )

        if DrillholeModel and drillholes:
            ct = ContentType.objects.get_for_model(DrillholeModel)
            ids = [str(r['id']) for r in drillholes]
            linked_doc_ids.update(
                DocLinkModel.objects.filter(content_type=ct, object_id__in=ids)
                .values_list('document_id', flat=True)
            )

        if processes:
            ct = ContentType.objects.get_for_model(Process)
            ids = [str(r['id']) for r in processes]
            linked_doc_ids.update(
                DocLinkModel.objects.filter(content_type=ct, object_id__in=ids)
                .values_list('document_id', flat=True)
            )

        if linked_doc_ids:
            documents = list(
                Document.objects.filter(org_filter, id__in=linked_doc_ids)
                .values('id', 'title', 'doc_type', 'created_at')[:20]
            )

    processes_list  = list(processes[:20])
    tenements_list  = list(tenements[:20])
    prospects_list  = list(prospects[:20])
    drillholes_list = list(drillholes[:20])

    return render(request, "core/partials/spatial_search_results.html", {
        "processes":  processes_list,
        "tenements":  tenements_list,
        "prospects":  prospects_list,
        "drillholes": drillholes_list,
        "documents":  documents,
        "total": sum([
            len(processes_list), len(tenements_list),
            len(prospects_list), len(drillholes_list),
        ]),
    })


# ---------- AI Report Generation & Document Analysis Pages ----------

@login_required
def report_list_page(request):
    clearance_rank = {"PUBLIC": 0, "INTERNAL": 1, "CONFIDENTIAL": 2, "JORC_APPROVED": 3}
    user_clearance = _get_clearance_level(request)
    user_rank = clearance_rank.get(user_clearance, 0)

    accessible_levels = [lvl for lvl, rank in clearance_rank.items() if rank <= user_rank]
    org_filter = _org_qs_filter(request)
    q = request.GET.get("q", "").strip()

    base_qs = (
        SavedReport.objects
        .filter(org_filter, clearance_level__in=accessible_levels)
        .select_related("process")
    )

    if q:
        sq = SearchQuery(q, search_type="websearch")
        recent_reports = (
            base_qs
            .annotate(rank=SearchRank("search_tsv", sq))
            .filter(Q(search_tsv=sq) | Q(title__icontains=q))
            .order_by("-rank", "-created_at")[:20]
        )
    else:
        recent_reports = base_qs.order_by("-created_at")[:20]

    recent_projects = Process.objects.filter(org_filter).order_by("-created_at")[:20]
    all_documents = Document.objects.filter(org_filter, is_latest=True).select_related("process").order_by("-created_at")

    return render(request, "core/report_list.html", {
        "recent_reports":  recent_reports,
        "recent_projects": recent_projects,
        "all_documents":   all_documents,
        "q":               q,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.GEOLOGIST_MINE,
    UserProfile.RoleChoices.METALLURGIST,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
def generate_report(request):
    """
    POST: generate (or retrieve cached) report for a process and redirect to the editor.
    GET:  redirect back to the report list.
    """
    if request.method != "POST":
        return redirect("report_list")

    process_id   = request.POST.get("process_id", "").strip()
    report_title = request.POST.get("report_title", "").strip()

    if not process_id:
        messages.error(request, "No project selected.")
        return redirect("report_list")

    org_filter = _org_qs_filter(request)
    if not Process.objects.filter(org_filter, pk=process_id).exists():
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    try:
        process = get_object_or_404(Process, org_filter, pk=process_id)
        bundle = _get_cached_report_bundle(process_id, clearance_level)
        md = bundle["md"]
        doc_ids = bundle.get("doc_ids", [])
    except Exception as e:
        log.error("Report generation failed during generate_report: %s", e)
        messages.error(request, f"Report generation failed: {e}")
        return redirect("report_list")

    import hashlib
    title = report_title or f"{process.name or 'Project'} Report"
    existing = SavedReport.objects.filter(
        process=process, title=title
    ).order_by("-version_number").first()

    if existing:
        report = SavedReport.create_version(
            parent=existing,
            content_md=md,
            user=request.user,
            reason=SavedReport.ChangeReason.REGENERATED,
        )
    else:
        report = SavedReport.objects.create(
            process=process,
            organisation=process.organisation,
            title=title,
            content_md=md,
            content_hash=hashlib.sha256(md.encode()).hexdigest(),
            created_by=request.user,
            version_number=1,
            change_reason=SavedReport.ChangeReason.GENERATED,
        )

    if doc_ids:
        report.source_documents.set(Document.objects.filter(pk__in=doc_ids))

    return redirect(reverse("report_editor", kwargs={"process_id": process_id}))


@login_required
def report_editor(request, process_id):
    """
    Serve the report editor page for a process
    loads markdown from cache (generated by generate_report before it redirects into here)
    """
    org_filter = _org_qs_filter(request)
    try:
        process = Process.objects.filter(org_filter).select_related("organisation").get(pk=process_id)
    except Process.DoesNotExist:
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    try:
        md = _get_cached_report_md(str(process_id), clearance_level)
    except Exception as e:
        log.error("Report editor cache miss for process %s: %s", process_id, e)
        md = f"# {process.name or 'Project'} Report\n\nReport generation failed: {e}"

    custom_title  = request.GET.get("title", "").strip()
    default_title = custom_title or f"{process.name or 'Project'} Report"

    return render(request, "core/report_editor.html", {
        "process": process,
        "markdown_content": md,
        "default_title": default_title,
        "saved_report": None,
        "save_url": reverse("save_report"),
        "export_url": reverse("export_report"),
    })


@login_required
def saved_report_editor(request, report_id):
    """ serve the report editor page for an existing saved report """
    report = get_object_or_404(
        SavedReport.objects.select_related("process", "organisation"),
        pk=report_id,
    )

    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and report.organisation
            and report.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    user_clearance = _get_clearance_level(request)
    clearance_rank = {"PUBLIC": 0, "INTERNAL": 1, "CONFIDENTIAL": 2, "JORC_APPROVED": 3}
    if clearance_rank.get(user_clearance, 0) < clearance_rank.get(report.clearance_level, 1):
        raise PermissionDenied

    process_prospects = (
        Prospect.objects.filter(process=report.process).order_by("name")
        if report.process else []
    )

    return render(request, "core/report_editor.html", {
        "process": report.process,
        "markdown_content": report.content_md,
        "default_title": report.title,
        "saved_report": report,
        "save_url": reverse("update_saved_report", kwargs={"report_id": report_id}),
        "export_url": reverse("export_report"),
        "process_prospects": process_prospects,
    })


@login_required
@require_POST
def save_report(request):
    """
    create a new SavedReport record from the user/editors current content
    returns JSON: {success: true, report_id: "...", redirect_url: "..."}
    """
    process_id = request.POST.get("process_id", "").strip()
    title      = request.POST.get("title", "").strip()
    content_md = request.POST.get("content_md", "").strip()

    if not title:
        return JsonResponse({"success": False, "error": "Title is required."}, status=400)
    if not content_md:
        return JsonResponse({"success": False, "error": "Report content is empty."}, status=400)

    process = None
    organisation = None
    if process_id:
        try:
            process = Process.objects.select_related("organisation").get(pk=process_id)
            organisation = process.organisation
        except Process.DoesNotExist:
            pass

    clearance_level = _get_clearance_level(request)
    created_by = request.user if request.user.is_authenticated else None

    existing = SavedReport.objects.filter(
        process=process, title=title
    ).order_by("-version_number").first()

    if existing:
        report = SavedReport.create_version(
            parent=existing,
            content_md=content_md,
            user=created_by,
            reason=SavedReport.ChangeReason.MANUAL_EDIT,
        )
        # Inherit source_documents from parent when no new ones available
        if not report.source_documents.exists() and existing.source_documents.exists():
            report.source_documents.set(existing.source_documents.all())
    else:
        import hashlib
        report = SavedReport.objects.create(
            process=process,
            organisation=organisation,
            title=title,
            content_md=content_md,
            content_hash=hashlib.sha256(content_md.encode()).hexdigest(),
            clearance_level=clearance_level,
            created_by=created_by,
            version_number=1,
            change_reason=SavedReport.ChangeReason.GENERATED,
        )

    # Populate source_documents from cached generation context if not already set
    if process_id and not report.source_documents.exists():
        try:
            cached = _get_cached_report_bundle(process_id, clearance_level)
            doc_ids = cached.get("doc_ids", [])
            if doc_ids:
                report.source_documents.set(Document.objects.filter(pk__in=doc_ids))
        except Exception:
            pass

    log_audit(request.user, AuditLog.ActionType.CREATE, report,
              f"Saved report '{title}' v{report.version_number}",
              ip_address=request.META.get("REMOTE_ADDR"))
    return JsonResponse({
        "success": True,
        "report_id": str(report.id),
        "version_number": report.version_number,
        "redirect_url": reverse("saved_report_editor", kwargs={"report_id": report.id}),
    })


@login_required
@require_POST
def update_saved_report(request, report_id):
    """
    overwrite of an existing SavedReport title and content
    Returns JSON: {success: true}
    """
    report = get_object_or_404(SavedReport, pk=report_id)

    if report.status in (SavedReport.Status.APPROVED, SavedReport.Status.PUBLISHED):
        return JsonResponse(
            {"success": False, "error": "This report is locked and cannot be edited."},
            status=403
        )

    is_admin = hasattr(request.user, "profile") and request.user.profile.role == "ADMIN"
    if report.created_by != request.user and not is_admin:
        return JsonResponse({"success": False, "error": "Permission denied."}, status=403)

    title      = request.POST.get("title", "").strip()
    content_md = request.POST.get("content_md", "").strip()

    if not title:
        return JsonResponse({"success": False, "error": "Title is required."}, status=400)
    if not content_md:
        return JsonResponse({"success": False, "error": "Report content is empty."}, status=400)

    new_version = SavedReport.create_version(
        parent=report,
        content_md=content_md,
        user=request.user,
        reason=SavedReport.ChangeReason.MANUAL_EDIT,
    )
    # Propagate source_documents to new version (new_version may equal report if content unchanged)
    if new_version.pk != report.pk and report.source_documents.exists():
        new_version.source_documents.set(report.source_documents.all())
    log_audit(request.user, AuditLog.ActionType.EDIT, new_version,
              f"Updated report '{new_version.title}' to v{new_version.version_number}",
              ip_address=request.META.get("REMOTE_ADDR"))
    return JsonResponse({
        "success": True,
        "new_version_id": str(new_version.id),
        "version_number": new_version.version_number,
    })


# ---------- JORC Approval Workflow Views ----------

@login_required
@require_POST
def submit_report_for_review(request, report_id):
    report = get_object_or_404(SavedReport, pk=report_id)
    if report.status != SavedReport.Status.DRAFT:
        messages.error(request, "Only draft reports can be submitted for review.")
        return redirect("saved_report_editor", report_id=report_id)

    workflow_type_raw = request.POST.get("workflow_type", "JORC").upper()
    valid_types = {t.value for t in ApprovalWorkflow.WorkflowType}
    workflow_type = workflow_type_raw if workflow_type_raw in valid_types else ApprovalWorkflow.WorkflowType.JORC
    submission_notes = request.POST.get("submission_notes", "").strip()

    workflow = ApprovalWorkflow.objects.create(
        content_type=ContentType.objects.get_for_model(SavedReport),
        object_id=report.pk,
        workflow_type=workflow_type,
        status=ApprovalWorkflow.Status.PENDING,
        submitted_by=request.user,
        submission_notes=submission_notes,
    )
    report.status = SavedReport.Status.UNDER_REVIEW
    report.approval_workflow = workflow
    report.save(update_fields=["status", "approval_workflow"])
    log_audit(request.user, AuditLog.ActionType.EDIT, report,
              f"Submitted for {workflow_type} review", ip_address=request.META.get("REMOTE_ADDR"))
    messages.success(request, f"Report submitted for {workflow.get_workflow_type_display()} review.")
    return redirect("saved_report_editor", report_id=report_id)


@login_required
@require_POST
def approve_report(request, report_id):
    from django.utils import timezone
    report = get_object_or_404(SavedReport, pk=report_id)
    profile = getattr(request.user, "profile", None)

    # Check permission based on the workflow type linked to this report
    workflow = getattr(report, "approval_workflow", None)
    wf_type = workflow.workflow_type if workflow else ApprovalWorkflow.WorkflowType.JORC
    if wf_type == ApprovalWorkflow.WorkflowType.VALMIN:
        can_approve = profile and profile.can_approve_valmin
    else:
        can_approve = profile and (
            profile.can_approve_jorc
            or profile.role == UserProfile.RoleChoices.COMPETENT_PERSON
        )
    if not can_approve:
        raise PermissionDenied

    if report.status != SavedReport.Status.UNDER_REVIEW:
        messages.error(request, "Only reports under review can be approved.")
        return redirect("saved_report_editor", report_id=report_id)

    approval_notes = request.POST.get("approval_notes", "").strip()
    report.status = SavedReport.Status.APPROVED
    report.save(update_fields=["status"])

    if workflow:
        workflow.status = ApprovalWorkflow.Status.APPROVED
        workflow.approved_by = request.user
        workflow.approval_notes = approval_notes
        workflow.reviewed_at = timezone.now()
        workflow.save(update_fields=["status", "approved_by", "approval_notes", "reviewed_at"])

    log_audit(request.user, AuditLog.ActionType.APPROVE, report,
              f"Approved ({wf_type})", ip_address=request.META.get("REMOTE_ADDR"))
    messages.success(request, "Report approved.")
    return redirect("saved_report_editor", report_id=report_id)


@login_required
@require_POST
def reject_report(request, report_id):
    from django.utils import timezone
    report = get_object_or_404(SavedReport, pk=report_id)
    profile = getattr(request.user, "profile", None)

    workflow = getattr(report, "approval_workflow", None)
    wf_type = workflow.workflow_type if workflow else ApprovalWorkflow.WorkflowType.JORC
    if wf_type == ApprovalWorkflow.WorkflowType.VALMIN:
        can_approve = profile and profile.can_approve_valmin
    else:
        can_approve = profile and (
            profile.can_approve_jorc
            or profile.role == UserProfile.RoleChoices.COMPETENT_PERSON
        )
    if not can_approve:
        raise PermissionDenied

    if report.status != SavedReport.Status.UNDER_REVIEW:
        messages.error(request, "Only reports under review can be rejected.")
        return redirect("saved_report_editor", report_id=report_id)

    approval_notes = request.POST.get("approval_notes", "").strip()
    report.status = SavedReport.Status.DRAFT
    report.approval_workflow = None
    report.save(update_fields=["status", "approval_workflow"])

    if workflow:
        workflow.status = ApprovalWorkflow.Status.REJECTED
        workflow.approved_by = request.user
        workflow.approval_notes = approval_notes
        workflow.reviewed_at = timezone.now()
        workflow.save(update_fields=["status", "approved_by", "approval_notes", "reviewed_at"])

    log_audit(request.user, AuditLog.ActionType.REJECT, report,
              f"Rejected ({wf_type}) — returned to draft", ip_address=request.META.get("REMOTE_ADDR"))
    messages.success(request, "Report returned to draft.")
    return redirect("saved_report_editor", report_id=report_id)


@login_required
@require_POST
def publish_report(request, report_id):
    report = get_object_or_404(SavedReport, pk=report_id)
    profile = getattr(request.user, "profile", None)
    if not (profile and (
        profile.can_approve_jorc
        or profile.role in (UserProfile.RoleChoices.ADMIN, UserProfile.RoleChoices.COMPETENT_PERSON)
    )):
        raise PermissionDenied
    if report.status != SavedReport.Status.APPROVED:
        messages.error(request, "Only approved reports can be published.")
        return redirect("saved_report_editor", report_id=report_id)
    report.status = SavedReport.Status.PUBLISHED
    report.save(update_fields=["status"])
    log_audit(request.user, AuditLog.ActionType.APPROVE, report,
              "Published", ip_address=request.META.get("REMOTE_ADDR"))
    messages.success(request, "Report published.")
    return redirect("saved_report_editor", report_id=report_id)


@login_required
def report_history(request, process_id):
    """Show all versions of reports for a process, grouped by title."""
    # Get the latest version of each distinct report title
    org_filter = _org_qs_filter(request)
    reports = (
        SavedReport.objects
        .filter(org_filter, process_id=process_id)
        .order_by("title", "-version_number")
    )
    # Group by title to show each report with its version chain
    from itertools import groupby
    grouped = {
        title: list(versions)
        for title, versions in groupby(reports, key=lambda r: r.title)
    }
    return render(request, "core/report_history.html", {"grouped": grouped, "process_id": process_id})

@login_required
def report_version_detail(request, report_id):
    """View a specific report version."""
    report = get_object_or_404(SavedReport, pk=report_id)
    all_versions = SavedReport.objects.filter(
        process=report.process, title=report.title
    ).order_by("-version_number")

    log_audit(
        user=request.user,
        action=AuditLog.ActionType.VIEW,
        obj=report,
        description=f"Viewed '{report.title}' v{report.version_number}",
        ip_address=request.META.get("REMOTE_ADDR"),
        user_agent=request.META.get("HTTP_USER_AGENT", ""),
    )
    return render(request, "core/report_version_detail.html", {
        "report": report,
        "all_versions": all_versions,
    })

@login_required
@require_GET
def all_reports_history(request):
    """Show all saved reports grouped by project, with optional full-text search."""
    org_filter = _org_qs_filter(request)
    q = request.GET.get("q", "").strip()

    reports = (
        SavedReport.objects
        .filter(org_filter)
        .select_related("process", "created_by")
        .order_by("process__name", "title", "-version_number")
    )

    if q:
        sq = SearchQuery(q, search_type="websearch")
        reports = (
            reports
            .annotate(rank=SearchRank("search_tsv", sq))
            .filter(Q(search_tsv=sq) | Q(title__icontains=q))
            .order_by("-rank", "-created_at")
        )

    total = reports.count() if q else None

    grouped = {}
    for report in reports:
        project_name = report.process.name if report.process else "No Project"
        process_id = str(report.process.id) if report.process else None
        if project_name not in grouped:
            grouped[project_name] = {"process_id": process_id, "titles": {}}
        if report.title not in grouped[project_name]["titles"]:
            grouped[project_name]["titles"][report.title] = []
        grouped[project_name]["titles"][report.title].append(report)

    return render(request, "core/all_reports_history.html", {
        "grouped": grouped,
        "q": q,
        "total": total,
    })


@login_required
@require_POST
def export_report(request):
    """
    Export the current markdown content as PDF or DOCX.
    POST params:
        format     — "pdf" or "docx"
        content_md — the markdown string to render
        title      — used as the filename
        report_id  — (optional) UUID of a SavedReport; if provided, a Sources section is appended
    """
    fmt       = request.POST.get("format", "pdf").lower()
    md_text   = request.POST.get("content_md", "")
    title     = request.POST.get("title", "report")
    report_id = request.POST.get("report_id", "").strip()
    slug      = re.sub(r"[^\w-]", "_", title)

    # Build sources list if a saved report ID was provided
    source_docs = []
    if report_id:
        try:
            saved = SavedReport.objects.prefetch_related("source_documents").get(pk=report_id)
            source_docs = list(saved.source_documents.order_by("title"))
        except (SavedReport.DoesNotExist, Exception):
            pass

    def _render_pdf(md_text, source_docs):
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        h1     = ParagraphStyle("h1", parent=styles["Heading1"], textColor=colors.HexColor("#0e7490"), spaceAfter=10)
        h2     = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#155e75"), spaceAfter=6)
        h3     = ParagraphStyle("h3", parent=styles["Heading3"], textColor=colors.HexColor("#1e4d5c"), spaceAfter=4)
        body   = ParagraphStyle("body", parent=styles["Normal"], spaceAfter=6, leading=16)
        bullet = ParagraphStyle("bullet", parent=styles["Normal"], leftIndent=20, spaceAfter=4, bulletIndent=10, leading=16)

        story = []
        for line in md_text.splitlines():
            if line.startswith("### "):
                story.append(Paragraph(line[4:], h3))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], h1))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {line[2:]}", bullet))
            elif re.match(r"^\d+\. ", line):
                story.append(Paragraph(re.sub(r"^\d+\. ", "", line), bullet))
            elif line.strip() == "":
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(line, body))

        if source_docs:
            story.append(Spacer(1, 20))
            story.append(Paragraph("Sources", h2))
            story.append(Paragraph(
                "The following documents were used as context for this report:", body
            ))
            for i, d in enumerate(source_docs, 1):
                doc_date = d.timestamp.strftime("%Y-%m-%d") if d.timestamp else ""
                story.append(Paragraph(
                    f"{i}. {d.title} — {d.doc_type or 'Document'} — {doc_date} — {d.confidentiality or ''}",
                    bullet
                ))

        doc.build(story)
        buf.seek(0)
        return buf.getvalue()

    def _render_docx(md_text, source_docs):
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in md_text.splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        if source_docs:
            doc.add_page_break()
            doc.add_heading("Sources", level=2)
            doc.add_paragraph("The following documents were used as context for this report:")
            for i, d in enumerate(source_docs, 1):
                doc_date = d.timestamp.strftime("%Y-%m-%d") if d.timestamp else ""
                doc.add_paragraph(
                    f"{i}. {d.title} — {d.doc_type or 'Document'} — {doc_date} — {d.confidentiality or ''}",
                    style="List Number"
                )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    if fmt == "pdf":
        content = _render_pdf(md_text, source_docs)
        response = HttpResponse(content, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{slug}_report.pdf"'
        return response

    if fmt == "docx":
        content = _render_docx(md_text, source_docs)
        response = HttpResponse(
            content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{slug}_report.docx"'
        return response

    return JsonResponse({"error": "Invalid format. Use 'pdf' or 'docx'."}, status=400)


@login_required
def report_detail(request, report_id):
    return render(request, "core/report_detail.html", {
        "report_id": report_id,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.ADMIN,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
)
def audit_log_view(request):
    """Filterable audit trail for ADMIN/DATA_MANAGER/OPS_MANAGER users."""
    import csv
    from django.http import StreamingHttpResponse

    qs = AuditLog.objects.select_related("user", "content_type").order_by("-timestamp")
    org_filter = _org_qs_filter(request)

    action    = request.GET.get("action", "").strip()
    username  = request.GET.get("username", "").strip()
    date_from = request.GET.get("date_from", "").strip()
    date_to   = request.GET.get("date_to", "").strip()
    obj_type  = request.GET.get("obj_type", "").strip()

    if action:
        qs = qs.filter(action=action)
    if username:
        qs = qs.filter(user__username__icontains=username)
    if date_from:
        qs = qs.filter(timestamp__date__gte=date_from)
    if date_to:
        qs = qs.filter(timestamp__date__lte=date_to)
    if obj_type:
        qs = qs.filter(content_type__model=obj_type.lower())

    if request.GET.get("export") == "csv":
        def _rows():
            yield ["Timestamp", "User", "Action", "Object Type", "Object ID", "Description", "IP Address"]
            for entry in qs.iterator(chunk_size=500):
                yield [
                    entry.timestamp.strftime("%Y-%m-%d %H:%M:%S"),
                    entry.user.username if entry.user else "",
                    entry.action,
                    entry.content_type.model if entry.content_type else "",
                    str(entry.object_id),
                    entry.description,
                    entry.ip_address or "",
                ]

        class _EchoWriter:
            def write(self, value):
                return value

        writer = csv.writer(_EchoWriter())
        response = StreamingHttpResponse(
            (writer.writerow(row) for row in _rows()),
            content_type="text/csv",
        )
        response["Content-Disposition"] = 'attachment; filename="audit_log.csv"'
        return response

    page_obj = _paginate(qs, request, per_page=50)
    return render(request, "core/audit_log.html", {
        "page_obj": page_obj,
        "action_choices": AuditLog.ActionType.choices,
        "current_action": action,
        "current_username": username,
        "current_date_from": date_from,
        "current_date_to": date_to,
        "current_obj_type": obj_type,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.ADMIN,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
)
def approval_workflows_list(request):
    """List all ApprovalWorkflow records for ADMIN/approval-capable users."""
    org_filter = _org_qs_filter(request)

    # Get workflow IDs associated with SavedReports in this org
    report_ct = ContentType.objects.get_for_model(SavedReport)
    report_ids = SavedReport.objects.filter(org_filter).values_list("id", flat=True)

    qs = (
        ApprovalWorkflow.objects
        .filter(content_type=report_ct, object_id__in=report_ids)
        .select_related("submitted_by", "approved_by")
        .order_by("-submitted_at")
    )

    status_filter = request.GET.get("status")
    wf_type_filter = request.GET.get("workflow_type")
    if status_filter:
        qs = qs.filter(status=status_filter)
    if wf_type_filter:
        qs = qs.filter(workflow_type=wf_type_filter)

    page_obj = _paginate(qs, request, per_page=25)
    return render(request, "core/approval_workflows_list.html", {
        "page_obj": page_obj,
        "status_choices": ApprovalWorkflow.Status.choices,
        "workflow_type_choices": ApprovalWorkflow.WorkflowType.choices,
        "current_status": status_filter or "",
        "current_workflow_type": wf_type_filter or "",
    })


@login_required
def document_analysis_page(request):
    return render(request, "core/document_analysis.html", {
        "recent_docs": Document.objects.filter(_org_qs_filter(request)).select_related("process").order_by("-created_at"),
    })


@login_required
def analyze_document(request, pk):
    document = get_object_or_404(Document, pk=pk)

    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and document.organisation
            and document.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    text = (document.extracted_text or "").strip()
    if not text:
        messages.error(request, "This document has no extracted text to analyse.")
        return redirect("document_analysis_page")

    try:
        client = GraniteClient()

        prompt = f"""
You are analysing a mining/exploration document.

Provide:
- A short summary
- Key insights
- Risks or issues
- Important findings
- Suggested next steps

Return the analysis in this exact format:

## Summary
...

## Key Insights
- ...

## Risks
- ...

## Recommended Actions
- ...
Document title: {document.title}

Document text:
{text[:12000]}
"""

        analysis_text = client.complete(prompt)

        document.analysis_text = analysis_text
        document.save(update_fields=["analysis_text"])

        messages.success(request, "Analysis complete.")
        return redirect("document_analysis_detail", pk=document.pk)

    except Exception as e:
        messages.error(request, f"Analysis failed: {e}")
        return redirect("document_analysis_page")

@login_required
@require_GET
def export_document_analysis(request, pk):
    document = get_object_or_404(Document, pk=pk)

    md_text = (document.analysis_text or "").strip()
    if not md_text:
        return JsonResponse({"error": "No analysis available to export."}, status=400)

    fmt = request.GET.get("format", "pdf").lower()
    title = f"{document.title} Analysis"
    slug = re.sub(r"[^\w-]", "_", title)

    if fmt == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2*cm,
            rightMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=colors.HexColor("#0e7490"), spaceAfter=10)
        h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#155e75"), spaceAfter=6)
        h3 = ParagraphStyle("h3", parent=styles["Heading3"], textColor=colors.HexColor("#1e4d5c"), spaceAfter=4)
        body = ParagraphStyle("body", parent=styles["Normal"], spaceAfter=6, leading=16)
        bullet = ParagraphStyle("bullet", parent=styles["Normal"], leftIndent=20, spaceAfter=4, bulletIndent=10, leading=16)

        story = []
        for line in md_text.splitlines():
            if line.startswith("### "):
                story.append(Paragraph(line[4:], h3))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], h1))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {line[2:]}", bullet))
            elif re.match(r"^\d+\. ", line):
                story.append(Paragraph(re.sub(r"^\d+\. ", "", line), bullet))
            elif line.strip() == "":
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(line, body))

        doc.build(story)
        buf.seek(0)
        response = HttpResponse(buf.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{slug}_analysis.pdf"'
        return response

    if fmt == "docx":
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in md_text.splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{slug}_analysis.docx"'
        return response

    return JsonResponse({"error": "Invalid format. Use 'pdf' or 'docx'."}, status=400)